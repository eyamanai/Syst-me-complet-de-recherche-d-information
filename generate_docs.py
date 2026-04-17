
import os

DOCS_TXT = [
    ("doc_txt_01_blockchain.txt", "Blockchain and Distributed Ledger Technology",
     """Blockchain is a distributed ledger technology that records transactions across multiple computers.
Each block contains a cryptographic hash of the previous block, forming a chain.
Bitcoin was the first application of blockchain technology, introduced in 2008.
Smart contracts are self-executing contracts with terms written directly into code.
Ethereum extended blockchain with programmable smart contracts and decentralized applications.
Consensus mechanisms like Proof of Work and Proof of Stake validate transactions.
Decentralized finance (DeFi) uses blockchain to recreate financial instruments without banks.
Non-fungible tokens (NFTs) represent unique digital assets on the blockchain."""),

    ("doc_txt_02_robotics.txt", "Robotics and Automation Systems",
     """Robotics combines mechanical engineering, electrical engineering, and computer science.
Industrial robots automate repetitive manufacturing tasks with high precision.
Autonomous mobile robots navigate environments using sensors and AI algorithms.
Computer vision enables robots to perceive and interpret their visual environment.
Collaborative robots (cobots) work safely alongside human workers.
Robot operating system (ROS) provides tools and libraries for robot software development.
Inverse kinematics calculates joint angles needed to reach a target position.
Swarm robotics coordinates large numbers of simple robots to achieve complex goals."""),

    ("doc_txt_03_quantum.txt", "Quantum Computing Fundamentals",
     """Quantum computing harnesses quantum mechanical phenomena to perform calculations.
Qubits can exist in superposition, representing 0 and 1 simultaneously.
Quantum entanglement allows correlations between qubits regardless of distance.
Quantum gates manipulate qubits to perform computational operations.
Shor's algorithm can factor large numbers exponentially faster than classical computers.
Grover's algorithm provides quadratic speedup for searching unsorted databases.
Quantum error correction addresses the challenge of qubit decoherence.
Near-term quantum devices are called Noisy Intermediate-Scale Quantum computers."""),

    ("doc_txt_04_iot.txt", "Internet of Things and Smart Devices",
     """The Internet of Things connects physical devices to the internet to collect and share data.
Smart home devices include thermostats, security cameras, and voice assistants.
Industrial IoT optimizes manufacturing and supply chain processes through sensor data.
Edge computing processes IoT data locally to reduce latency and bandwidth usage.
MQTT is a lightweight messaging protocol commonly used in IoT applications.
IoT security is critical as billions of connected devices create attack surfaces.
Wearable devices monitor health metrics like heart rate and sleep patterns.
5G networks provide the bandwidth and low latency needed for IoT deployment."""),

    ("doc_txt_05_devops.txt", "DevOps Culture and Practices",
     """DevOps combines software development and IT operations to shorten delivery cycles.
Continuous integration merges code changes frequently to detect problems early.
Continuous deployment automates the release of software to production environments.
Infrastructure as code manages IT infrastructure using configuration files.
Monitoring and observability provide insight into system performance and behavior.
Containerization with Docker packages applications with their dependencies.
Kubernetes orchestrates containers at scale across server clusters.
Site reliability engineering applies software engineering practices to operations."""),

    ("doc_txt_06_ethics_ai.txt", "Ethics in Artificial Intelligence",
     """AI ethics addresses the moral implications of artificial intelligence systems.
Algorithmic bias occurs when AI systems produce unfair outcomes for certain groups.
Transparency requires AI systems to explain their decisions to affected users.
Privacy concerns arise from AI systems collecting and processing personal data.
Accountability determines who is responsible when AI systems cause harm.
Fairness in machine learning ensures equitable treatment across demographic groups.
The alignment problem addresses how to ensure AI pursues intended human values.
Regulatory frameworks for AI are being developed by governments worldwide."""),

    ("doc_txt_07_bioinformatics.txt", "Bioinformatics and Computational Biology",
     """Bioinformatics applies computational methods to analyze biological data.
Sequence alignment compares DNA, RNA, or protein sequences to find similarities.
Genome assembly reconstructs complete genomic sequences from short reads.
Gene expression analysis measures which genes are active in different conditions.
Protein structure prediction determines the 3D shape of proteins from sequences.
Phylogenetic analysis reconstructs evolutionary relationships between organisms.
BLAST searches biological sequence databases for similar sequences quickly.
Machine learning increasingly drives drug discovery and genomic research."""),

    ("doc_txt_08_cybersecurity_advanced.txt", "Advanced Cybersecurity Techniques",
     """Penetration testing simulates attacks to identify security vulnerabilities proactively.
Zero-day exploits target previously unknown software vulnerabilities.
Social engineering manipulates humans rather than attacking technical systems.
Advanced persistent threats involve prolonged, targeted attacks by sophisticated actors.
Threat intelligence collects and analyzes data about cybersecurity threats.
Security information and event management systems correlate security logs.
Zero trust architecture assumes no user or device is inherently trustworthy.
Homomorphic encryption enables computation on encrypted data without decryption."""),

    ("doc_txt_09_augmented_reality.txt", "Augmented and Virtual Reality",
     """Augmented reality overlays digital information onto the real world environment.
Virtual reality creates fully immersive digital environments for users.
Mixed reality blends physical and digital worlds in real time interaction.
Head-mounted displays provide immersive visual experiences for VR and AR.
Spatial computing enables interaction with three-dimensional digital content.
SLAM technology simultaneously maps environments and tracks device position.
Haptic feedback provides tactile sensations to enhance virtual reality immersion.
Enterprise applications of AR include training, maintenance, and remote assistance."""),

    ("doc_txt_10_recommender.txt", "Recommender Systems and Personalization",
     """Recommender systems suggest items likely to interest individual users.
Collaborative filtering identifies users with similar preferences for recommendations.
Content-based filtering recommends items similar to those a user previously liked.
Matrix factorization decomposes user-item interaction matrices for recommendations.
The cold start problem occurs when systems lack data for new users or items.
Implicit feedback from user behavior like clicks provides recommendation signals.
Deep learning models capture complex patterns in user preferences and item features.
Evaluation metrics include precision, recall, and normalized discounted cumulative gain."""),
]

DOCS_HTML = [
    ("doc_html_01_ml_basics.html", "Introduction to Machine Learning",
     """<html><head><title>Machine Learning Basics</title></head><body>
<h1>Introduction to Machine Learning</h1>
<p>Machine learning enables computers to learn from data without explicit programming.
The field has three main paradigms: supervised, unsupervised, and reinforcement learning.</p>
<h2>Supervised Learning</h2>
<p>Supervised learning uses labeled training data to learn mappings from inputs to outputs.
Common algorithms include linear regression, logistic regression, decision trees, and neural networks.
Cross-validation helps estimate model performance on unseen data and prevents overfitting.</p>
<h2>Model Evaluation</h2>
<p>Performance metrics depend on the task: accuracy, precision, recall for classification;
mean squared error, R-squared for regression. The bias-variance tradeoff is fundamental
to understanding model performance. Regularization techniques like L1 and L2 prevent overfitting.</p>
</body></html>"""),

    ("doc_html_02_databases.html", "Database Systems Overview",
     """<html><head><title>Database Systems</title></head><body>
<h1>Database Management Systems</h1>
<p>Databases organize and store data for efficient retrieval and management.
Relational databases use tables, SQL, and ACID properties for reliable transactions.</p>
<h2>NoSQL Databases</h2>
<p>NoSQL databases sacrifice some ACID properties for scalability and flexibility.
Document stores like MongoDB store data as JSON-like documents.
Key-value stores like Redis provide ultra-fast in-memory data access.
Graph databases like Neo4j model relationships between entities efficiently.</p>
<h2>Query Optimization</h2>
<p>Query optimizers find efficient execution plans for database queries.
Indexes speed up data retrieval at the cost of storage and write performance.
Partitioning distributes large tables across multiple storage units for scalability.</p>
</body></html>"""),

    ("doc_html_03_networks.html", "Computer Networking Fundamentals",
     """<html><head><title>Computer Networks</title></head><body>
<h1>Computer Networking</h1>
<p>Networks connect computers to share resources and communicate.
The Internet uses the TCP/IP protocol suite developed in the 1970s.</p>
<h2>Network Protocols</h2>
<p>HTTP and HTTPS enable communication on the World Wide Web.
DNS translates human-readable domain names to IP addresses.
SMTP, IMAP, and POP3 handle email transmission and retrieval.
FTP and SFTP transfer files between computers securely.</p>
<h2>Network Security</h2>
<p>Firewalls filter network traffic based on security rules.
VPNs create encrypted tunnels for secure remote access.
TLS encryption protects data in transit from eavesdropping.
Intrusion detection systems monitor networks for suspicious activity.</p>
</body></html>"""),

    ("doc_html_04_python.html", "Python Programming Guide",
     """<html><head><title>Python Programming</title></head><body>
<h1>Python Programming Language</h1>
<p>Python is known for its readability and versatility across many domains.
It supports procedural, object-oriented, and functional programming paradigms.</p>
<h2>Key Features</h2>
<p>Python's dynamic typing allows variables to hold any type of value.
List comprehensions create new lists concisely: [x*2 for x in range(10)].
Generators use yield to produce values lazily, saving memory for large datasets.
Decorators wrap functions to add behavior without modifying the original code.</p>
<h2>Scientific Computing</h2>
<p>NumPy provides efficient array operations for numerical computation.
SciPy builds on NumPy with algorithms for optimization, statistics, and signal processing.
Matplotlib creates publication-quality visualizations from data.
Scikit-learn provides machine learning algorithms with a consistent API.</p>
</body></html>"""),

    ("doc_html_05_cloud.html", "Cloud Computing Services",
     """<html><head><title>Cloud Computing</title></head><body>
<h1>Cloud Computing</h1>
<p>Cloud computing delivers on-demand computing resources over the internet.
The major providers are Amazon Web Services, Microsoft Azure, and Google Cloud.</p>
<h2>Service Models</h2>
<p>Infrastructure as a Service provides virtualized computing resources on demand.
Platform as a Service offers development environments and tools in the cloud.
Software as a Service delivers complete applications over the internet.
Serverless computing runs code in response to events without managing servers.</p>
<h2>Cloud Storage</h2>
<p>Object storage like S3 stores unstructured data at massive scale cheaply.
Block storage provides low-latency storage for databases and applications.
File storage offers shared network file systems for multiple instances.
Content delivery networks cache content geographically for faster delivery.</p>
</body></html>"""),
]

DOCS_DOCX_TEXT = [
    ("doc_docx_01_software_arch.txt", "Software Architecture Patterns",
     """Software architecture defines the high-level structure of software systems.
Microservices architecture decomposes applications into independently deployable services.
Event-driven architecture uses events to trigger and communicate between services.
The Model-View-Controller pattern separates concerns in user interface applications.
Repository pattern abstracts data access logic from business logic in applications.
CQRS separates command (write) and query (read) responsibilities for scalability.
Hexagonal architecture isolates application core from external systems and infrastructure.
Domain-driven design aligns software structure with business domain concepts."""),

    ("doc_docx_02_data_engineering.txt", "Data Engineering and Pipelines",
     """Data engineering builds systems that collect, process, and store data at scale.
ETL pipelines extract data from sources, transform it, and load it into destinations.
Apache Spark provides distributed data processing for large-scale batch workloads.
Apache Kafka enables real-time streaming data pipelines between systems.
Data lakes store raw data in native format for future processing and analysis.
Data warehouses store structured, processed data optimized for analytical queries.
dbt transforms data within the warehouse using SQL and software engineering practices.
Data quality monitoring detects anomalies and ensures data meets defined standards."""),

    ("doc_docx_03_nlp_advanced.txt", "Advanced Natural Language Processing",
     """Large language models are trained on massive text corpora using self-supervised learning.
Attention mechanisms allow models to focus on relevant parts of input sequences.
The transformer architecture revolutionized NLP with parallel sequence processing.
BERT learns bidirectional representations by predicting masked words in context.
GPT models generate coherent text autoregressively given a starting prompt.
Fine-tuning adapts pre-trained language models to specific downstream tasks.
Retrieval-augmented generation combines language models with external knowledge bases.
Chain-of-thought prompting improves reasoning by eliciting step-by-step explanations."""),

    ("doc_docx_04_statistics.txt", "Statistics and Probability for Data Science",
     """Statistics provides tools for collecting, analyzing, and interpreting data.
Descriptive statistics summarize data using measures of central tendency and spread.
Inferential statistics draws conclusions about populations from sample data.
Hypothesis testing determines whether observed effects are statistically significant.
The p-value measures the probability of observing results as extreme by chance.
Bayesian inference updates probability estimates as new evidence becomes available.
Regression analysis models relationships between dependent and independent variables.
A/B testing compares two versions of something to determine which performs better."""),

    ("doc_docx_05_computer_graphics.txt", "Computer Graphics and Visualization",
     """Computer graphics creates and manipulates visual images using computers.
Rasterization converts geometric primitives into pixel arrays for display.
Ray tracing simulates the physical behavior of light for photorealistic rendering.
Shaders are programs that run on the GPU to determine pixel colors and effects.
3D modeling creates geometric representations of objects in three-dimensional space.
Animation systems simulate motion through keyframes or physics-based simulation.
OpenGL and Vulkan provide APIs for hardware-accelerated graphics rendering.
Texture mapping applies 2D images to 3D surfaces for realistic appearance."""),
]

PDF_EXTRA = [
    ("doc_16_deep_learning.pdf", "Deep Learning Architectures",
     """Deep learning uses multiple layers of neural networks to learn complex representations.
Convolutional neural networks excel at processing grid-structured data like images.
Recurrent neural networks process sequential data by maintaining hidden state.
Long short-term memory networks solve the vanishing gradient problem in RNNs.
Generative adversarial networks train two competing networks to generate realistic data.
Autoencoders learn compressed representations by encoding and then decoding inputs.
Batch normalization stabilizes training by normalizing layer activations.
Dropout regularization prevents overfitting by randomly disabling neurons during training."""),

    ("doc_17_reinforcement_learning.pdf", "Reinforcement Learning",
     """Reinforcement learning trains agents to take actions maximizing cumulative reward.
The agent interacts with an environment, receiving observations and rewards.
Q-learning learns action-value functions without a model of the environment.
Deep Q-networks combine Q-learning with deep neural networks for complex tasks.
Policy gradient methods directly optimize the policy for selecting actions.
Actor-critic methods combine value functions and policy gradient approaches.
Model-based reinforcement learning uses learned environment models for planning.
Multi-agent reinforcement learning studies the behavior of multiple interacting agents."""),

    ("doc_18_data_visualization.pdf", "Data Visualization Principles",
     """Data visualization communicates information through graphical representations.
Effective visualizations reveal patterns and trends not apparent in raw data.
Chart selection depends on the type of data and the message to convey.
Bar charts compare discrete categories using rectangular bars of varying heights.
Line charts show trends and changes over continuous variables like time.
Scatter plots reveal relationships and correlations between two continuous variables.
Heatmaps display matrix data using color intensity to encode values.
Interactive visualizations allow users to explore data through filtering and zooming."""),

    ("doc_19_parallel_computing.pdf", "Parallel and Distributed Computing",
     """Parallel computing uses multiple processors to solve problems simultaneously.
Flynn's taxonomy classifies parallel architectures by instruction and data streams.
Shared memory parallelism allows processors to access a common memory space.
Message passing interface enables communication between distributed memory processes.
GPU computing uses thousands of cores for highly parallel numerical computation.
MapReduce programming model processes large datasets across distributed clusters.
Load balancing distributes computational work evenly across processing units.
Amdahl's law limits the speedup achievable through parallelization of a program."""),

    ("doc_20_api_design.pdf", "API Design and Web Services",
     """APIs enable communication between different software applications and services.
RESTful APIs use HTTP methods and resources to define a uniform interface.
GraphQL allows clients to request exactly the data they need from APIs.
API versioning manages changes while maintaining backward compatibility for clients.
Authentication mechanisms like OAuth 2.0 secure access to API resources.
Rate limiting prevents API abuse by restricting the number of requests per period.
API documentation tools like Swagger generate interactive documentation from code.
Webhooks enable APIs to push notifications to clients when events occur."""),
]


def generer_txt(dossier):
    """Génère les fichiers TXT."""
    d = os.path.join(dossier, "txt")
    os.makedirs(d, exist_ok=True)
    for fname, title, content in DOCS_TXT:
        with open(os.path.join(d, fname), "w", encoding="utf-8") as f:
            f.write(f"{title}\n{'='*60}\n\n{content.strip()}\n")
    print(f"[OK] {len(DOCS_TXT)} fichiers TXT générés dans {d}/")


def generer_html(dossier):
    """Génère les fichiers HTML."""
    d = os.path.join(dossier, "html")
    os.makedirs(d, exist_ok=True)
    for fname, title, content in DOCS_HTML:
        with open(os.path.join(d, fname), "w", encoding="utf-8") as f:
            f.write(content)
    print(f"[OK] {len(DOCS_HTML)} fichiers HTML générés dans {d}/")


def generer_docx_txt(dossier):
    """Génère les fichiers DOCX (ou TXT en fallback)."""
    d = os.path.join(dossier, "docx")
    os.makedirs(d, exist_ok=True)
    try:
        import docx
        for fname, title, content in DOCS_DOCX_TEXT:
            doc = docx.Document()
            doc.add_heading(title, 0)
            for ligne in content.strip().split("\n"):
                if ligne.strip():
                    doc.add_paragraph(ligne.strip())
            out = os.path.join(d, fname.replace(".txt", ".docx"))
            doc.save(out)
        print(f"[OK] {len(DOCS_DOCX_TEXT)} fichiers DOCX générés dans {d}/")
    except ImportError:
        for fname, title, content in DOCS_DOCX_TEXT:
            with open(os.path.join(d, fname), "w", encoding="utf-8") as f:
                f.write(f"{title}\n{'='*60}\n\n{content.strip()}\n")
        print(f"[OK] {len(DOCS_DOCX_TEXT)} fichiers TXT (fallback DOCX) générés dans {d}/")


def generer_pdf_extra(dossier):
    """Génère les PDF supplémentaires (16-20)."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        for filename, title, content in PDF_EXTRA:
            path = os.path.join(dossier, "pdf", filename)
            c = canvas.Canvas(path, pagesize=A4)
            c.setFont("Helvetica-Bold", 14)
            c.drawString(50, 800, title)
            c.setFont("Helvetica", 10)
            y = 770
            for line in content.strip().split("\n"):
                line = line.strip()
                if not line:
                    y -= 6
                    continue
                while len(line) > 95:
                    c.drawString(50, y, line[:95])
                    line = line[95:]
                    y -= 14
                c.drawString(50, y, line)
                y -= 14
            c.save()
        print(f"[OK] {len(PDF_EXTRA)} PDF supplémentaires générés dans {dossier}/pdf/")
    except ImportError:
        print("[WARN] reportlab non installé, PDF supplémentaires ignorés.")


def generer_collection():
    dossier = "documents"
    os.makedirs(dossier, exist_ok=True)
    os.makedirs(os.path.join(dossier, "pdf"), exist_ok=True)

    # Générer les 15 PDF originaux dans documents/pdf/
    from generate_docs_original import generer_collection as gen_orig
    gen_orig()

    # Générer les formats supplémentaires
    generer_txt(dossier)
    generer_html(dossier)
    generer_docx_txt(dossier)
    generer_pdf_extra(dossier)

    total = 15 + len(DOCS_TXT) + len(DOCS_HTML) + len(DOCS_DOCX_TEXT) + len(PDF_EXTRA)
    print(f"\n[TOTAL] {total} documents générés dans '{dossier}/'")
    print("  - PDF      : 15 + 5 = 20")
    print(f"  - TXT      : {len(DOCS_TXT)}")
    print(f"  - HTML     : {len(DOCS_HTML)}")
    print(f"  - DOCX/TXT : {len(DOCS_DOCX_TEXT)}")


if __name__ == "__main__":
    # Génération simplifiée sans dépendance circulaire
    dossier = "documents"
    os.makedirs(dossier, exist_ok=True)
    generer_txt(dossier)
    generer_html(dossier)
    generer_docx_txt(dossier)
    print("\nPour les PDF, lancez d'abord : python generate_docs_pdf.py")
